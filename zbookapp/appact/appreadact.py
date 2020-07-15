from django.shortcuts import  HttpResponse
from django.db import connection
import json
from zbookapp import public
from zbookapp import models
from zbookapp.createanswer.Answermatch.matchfun.QAFun import alloperate
import datetime
import docx
import os
import time



#阅读处理
def admin(request):
    log = public.logger
    log.info('----------------------appreadact_begin---------------------------')

    if request.method == "POST":
        # #请求body转为json
        tmp=request.body
        tmp=tmp.decode(encoding='utf-8')
        request_body=json.loads(tmp)
        trantype = request_body['trantype']
        log.info('trantype=[%s]' % trantype)
        fun_name = trantype
        if globals().get(fun_name):
            resp = globals()[fun_name](request, request_body)
        else:
            s = public.setrespinfo({"respcode": "100000", "respmsg": "api error"})
            resp = HttpResponse(s)

    elif request.method == "GET":
        s = public.setrespinfo({"respcode": "000000", "respmsg": "api error"})
        resp = HttpResponse(s)

    log.info('----------------------appreadact_end---------------------------')
    return resp

# 文档状态设置
def updfilestatus(request, request_body):
    log = public.logger
    log.info('----------------------Admin-getbookinfo-begin---------------------------')
    jsondata = {
        "respcode": "000000",
        "respmsg": "操作成功",
        "trantype": request_body.get('trantype', None),
    }
    uid = request_body.get('userid', None)
    status = request_body.get('status', None)
    fileid = request_body.get('fileid', None)

    if uid:
        sql = "select * from zbookapp_user where user_id=%s"
        cur = connection.cursor()
        cur.execute(sql, uid)
        row = cur.fetchone()
        cur.close()
        if not row:
            s = public.setrespinfo({"respcode": "300004", "respmsg": "用户不存在！"})
            return s
    else:
        s = public.setrespinfo({"respcode": "300003", "respmsg": "请登录！"})
        return s
    upd_sql = "update zbookapp_bookfile set status=%s where file_id=%s and user_id=%s"
    cur = connection.cursor()
    cur.execute(upd_sql,(status,fileid,uid))
    connection.commit()
    s = json.dumps(jsondata, cls=models.JsonCustomEncoder, ensure_ascii=False)
    log.info('----------------------Admin-getbookinfo-end---------------------------')
    return HttpResponse(s)

#书架状态设置
def updshelfstatus(request, request_body):
    log = public.logger
    log.info('----------------------Admin-updshelfstatus-begin---------------------------')
    jsondata = {
        "respcode": "000000",
        "respmsg": "操作成功",
        "trantype": request_body.get('trantype', None),
    }
    uid = request_body.get('userid', None)
    status = request_body.get('status', None)
    fileid = request_body.get('fileid', None)

    if uid:
        sql = "select * from zbookapp_user where user_id=%s"
        cur = connection.cursor()
        cur.execute(sql, uid)
        row = cur.fetchone()
        cur.close()
        if not row:
            s = public.setrespinfo({"respcode": "300004", "respmsg": "用户不存在！"})
            return s
    else:
        s = public.setrespinfo({"respcode": "300003", "respmsg": "请登录！"})
        return s
    upd_sql = "update book_shelf set status=%s where file_id=%s and user_id=%s"
    cur = connection.cursor()
    cur.execute(upd_sql,(status,fileid,uid))
    connection.commit()
    s = json.dumps(jsondata, cls=models.JsonCustomEncoder, ensure_ascii=False)
    log.info('----------------------Admin-updshelfstatus-end---------------------------')
    return HttpResponse(s)

#收藏状态设置
def updcollstatus(request, request_body):
    log = public.logger
    log.info('----------------------Admin-updshelfstatus-begin---------------------------')
    jsondata = {
        "respcode": "000000",
        "respmsg": "操作成功",
        "trantype": request_body.get('trantype', None),
    }
    uid = request_body.get('userid', None)
    fileid = request_body.get('fileid', None)

    if uid:
        sql = "select * from zbookapp_user where user_id=%s"
        cur = connection.cursor()
        cur.execute(sql, uid)
        row = cur.fetchone()
        cur.close()
        if not row:
            s = public.setrespinfo({"respcode": "300004", "respmsg": "用户不存在！"})
            return s
    else:
        s = public.setrespinfo({"respcode": "300003", "respmsg": "请登录！"})
        return s
    del_sql = "delete from zbookapp_collection where file_id=%s and user_id=%s"
    cur = connection.cursor()
    cur.execute(del_sql,(fileid,uid))
    connection.commit()
    s = json.dumps(jsondata, cls=models.JsonCustomEncoder, ensure_ascii=False)
    log.info('----------------------Admin-updshelfstatus-end---------------------------')
    return HttpResponse(s)

#加入收藏
def joincol(request, request_body):
    log = public.logger
    log.info('----------------------Admin-getbookinfo-begin---------------------------')
    jsondata = {
        "respcode": "000000",
        "respmsg": "上传成功",
        "trantype": request_body.get('trantype', None),
    }
    uid = request_body.get('userid', None)
    fileid = request_body.get('fileid', None)
    coltext='收藏'

    if uid:
        sql = "select * from zbookapp_user where user_id=%s"
        cur = connection.cursor()
        cur.execute(sql, uid)
        row = cur.fetchone()
        cur.close()
        if not row:
            s = public.setrespinfo({"respcode": "300004", "respmsg": "用户不存在！"})
            return s
    else:
        s = public.setrespinfo({"respcode": "300003", "respmsg": "请登录！"})
        return s
    select_sql = "select * from zbookapp_collection where file_id=%s and user_id=%s and status='0' "
    cur = connection.cursor()
    cur.execute(select_sql,(fileid,uid))
    row = cur.fetchone()
    if row:
        s = public.setrespinfo({"respcode": "300006", "respmsg": "已加入收藏!"})
        return s

    nowTime = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    try:
        join_sql = "INSERT INTO zbookapp_collection (file_id, user_id, tran_date, status) VALUES (%s, %s, %s, %s)"
        cur = connection.cursor()
        cur.execute(join_sql, (fileid,uid,nowTime,'0'))

        cur.close()
        coltext = '已收藏'
    except Exception as e:
        print(e)


    jsondata['coltext'] = coltext
    s = json.dumps(jsondata, cls=models.JsonCustomEncoder, ensure_ascii=False)
    log.info('----------------------Admin-getbookinfo-end---------------------------')
    return HttpResponse(s)

#加入书架
def joinshelf(request, request_body):
    log = public.logger
    log.info('----------------------Admin-joinshelf-begin---------------------------')
    jsondata = {
        "respcode": "000000",
        "respmsg": "上传成功",
        "trantype": request_body.get('trantype', None),
    }
    uid = request_body.get('userid', None)
    fileid = request_body.get('fileid', None)
    shelftext='加入书架'

    if uid:
        sql = "select * from zbookapp_user where user_id=%s"
        cur = connection.cursor()
        cur.execute(sql, uid)
        row = cur.fetchone()
        cur.close()
        if not row:
            s = public.setrespinfo({"respcode": "300004", "respmsg": "用户不存在！"})
            return s
    else:
        s = public.setrespinfo({"respcode": "300003", "respmsg": "请登录！"})
        return s

    # 判断是否已经加入书架
    select_sql = "select * from book_shelf where user_id=%s and file_id=%s and status='0' "
    cur = connection.cursor()
    cur.execute(select_sql,(uid,fileid))
    row = cur.fetchone()
    if row:
        s = public.setrespinfo({"respcode": "300007", "respmsg": "已加入书架!"})
        return s

    try:
        join_sql = "INSERT INTO book_shelf(id, user_id, file_id, tran_date,upd_date, status) VALUES(%s, %s, %s, %s, %s, %s)"
        cur = connection.cursor()
        cur.execute(join_sql, (None,uid,fileid,datetime.datetime.now(),datetime.datetime.now(),'0'))
        shelftext = '已加入书架'
    except Exception as e:
        print(e)


    jsondata['shelftext'] = shelftext
    s = json.dumps(jsondata, cls=models.JsonCustomEncoder, ensure_ascii=False)
    log.info('----------------------Admin-joinshelf-end---------------------------')
    return HttpResponse(s)

# 智能阅读处理
def create_answer(request, request_body):
    log = public.logger
    log.info('----------------------Admin-create_answer-begin---------------------------')
    jsondata = {
        "respcode": "000000",
        "respmsg": "处理完成",
        "trantype": request_body.get('trantype', None),
    }

    question = request_body.get('question', None)
    fileid = request_body.get('fileid', None)
    if not fileid:
        s = public.setrespinfo({"respcode": "400001", "respmsg": "文件不存在！"})
        return s
    if len(question)==0:
        s = public.setrespinfo({"respcode": "400002", "respmsg": "请输入问题！"})
        return s
    # 查询文件信息
    file_sql = "select * from zbookapp_bookfile where file_id = %s "
    cur = connection.cursor()
    cur.execute(file_sql, fileid)
    row = cur.fetchone()
    if row:
        md5_filename = row[3]
        md5file = md5_filename.split('.')[0]
        data_path = public.localhome + "fileup//"+md5_filename
        # print('执行到这了')
        # # time.sleep(5)
        tmp_answer = []
        try:
            tmp_answer = alloperate(md5file,data_path,question)
        except Exception as e:
            print('e=',e)
        answer = []
        if tmp_answer:
            for index,item in enumerate(tmp_answer):
                tmp_lineid = str(item['line_id'])
                item['line_id'] = tmp_lineid
                item['label'] = '答案'+ str(index+1)
                answer.append(item)
        jsondata['answerlist'] = answer
    else:
        s = public.setrespinfo({"respcode": "400001", "respmsg": "文件不存在！"})
        return s



    s = json.dumps(jsondata, cls=models.JsonCustomEncoder, ensure_ascii=False)
    log.info('----------------------Admin-create_answer-end---------------------------')
    return HttpResponse(s)

# 文档预览
def docpreview(request,request_body):
    log = public.logger
    log.info('----------------------Admin-docpreview-begin---------------------------')
    jsondata = {
        "respcode": "000000",
        "respmsg": "处理完成",
        "trantype": request_body.get('trantype', None),
    }

    fileid = request_body.get('fileid', None)
    begin_num = request_body.get('begin_num',0)
    end_num = begin_num + 20
    if not fileid:
        s = public.setrespinfo({"respcode": "400001", "respmsg": "文档不存在！"})
        return s
    # 查询文件信息
    file_sql = "select * from zbookapp_bookfile where file_id = %s "
    cur = connection.cursor()
    cur.execute(file_sql, fileid)
    row = cur.fetchone()
    print('row=',row)
    if row:
        md5_filename = row[3]
        data_path = public.localhome + "fileup//" + md5_filename
        doc_data = read_data(data_path,begin_num,end_num)
        jsondata['doc_data'] = doc_data
        jsondata['begin_num'] = end_num+1
    else:
        s = public.setrespinfo({"respcode": "400001", "respmsg": "文档不存在！"})
        return s

    s = json.dumps(jsondata, cls=models.JsonCustomEncoder, ensure_ascii=False)
    log.info('----------------------Admin-docpreview-end---------------------------')
    return HttpResponse(s)

def read_data(data_path,begin_num=0,end_num=10):
    file_type = data_path.split('.')[-1]
    if file_type == 'txt':
        docs = open(data_path, encoding='utf-8').readlines()[begin_num:end_num]
        data = ''
        for item in docs:
            data = data + item.strip()
    elif file_type == 'docx':
        data = ''
        file = docx.Document(data_path)
        print('file=',file,type(file.paragraphs))

        for para in (file.paragraphs)[begin_num:end_num]:
            data = data + para.text+'\n'
    return data